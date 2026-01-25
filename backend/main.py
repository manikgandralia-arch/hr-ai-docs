import os
import json
from datetime import datetime

from dotenv import load_dotenv
from docx import Document
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

from groq import Groq


# ---------------- INIT ----------------

load_dotenv()

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
GENERATED_DIR = os.path.join(BASE_DIR, "generated")

os.makedirs(GENERATED_DIR, exist_ok=True)

groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

app = FastAPI()


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------- MODELS ----------------

class EmployeeData(BaseModel):
    employee_name: str
    employee_address: str
    designation: str
    department: str
    company_name: str
    work_location: str = ""
    reporting_manager: str = ""
    offer_date: str = ""
    joining_date: str = ""
    end_date: str = ""
    termination_date: str = ""
    reason: str = ""
    ctc: str = ""
    probation_months: str = ""
    notice_period: str = ""
    working_hours: str = ""
    hr_name: str


# ---------------- HELPERS ----------------

def replace_placeholders(doc: Document, mapping: dict):

    for para in doc.paragraphs:
        for key, value in mapping.items():
            if key in para.text:
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in mapping.items():
                        if key in para.text:
                            for run in para.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)


def ai_review_document(text: str, doc_type: str):

    prompt = f"""
You are an HR compliance assistant for India.
Check this {doc_type} for missing or risky clauses.

Return ONLY valid JSON with keys:
risk_score (1-10),
missing_clauses (array),
weak_clauses (array),
suggested_text (array)

Document:
{text}
"""

    try:
        resp = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {"role": "system", "content": "You output ONLY JSON."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
        )

        output = resp.choices[0].message.content.strip()

        return json.loads(output)

    except Exception as e:
        return {"error": str(e)}


def generate_document(data, template, prefix, doc_type):

    template_path = os.path.join(TEMPLATE_DIR, template)

    if not os.path.exists(template_path):
        raise HTTPException(status_code=500, detail="Template missing")

    doc = Document(template_path)

    mapping = {
        "{EMPLOYEE_NAME}": data.employee_name,
        "{EMPLOYEE_ADDRESS}": data.employee_address,
        "{DESIGNATION}": data.designation,
        "{DEPARTMENT}": data.department,
        "{COMPANY_NAME}": data.company_name,
        "{WORK_LOCATION}": data.work_location,
        "{REPORTING_MANAGER}": data.reporting_manager,
        "{OFFER_DATE}": data.offer_date,
        "{JOINING_DATE}": data.joining_date,
        "{END_DATE}": data.end_date,
        "{TERMINATION_DATE}": data.termination_date,
        "{REASON}": data.reason,
        "{CTC}": data.ctc,
        "{PROBATION_MONTHS}": data.probation_months,
        "{NOTICE_PERIOD}": data.notice_period,
        "{WORKING_HOURS}": data.working_hours,
        "{HR_NAME}": data.hr_name,
        "{DATE}": datetime.now().strftime("%Y-%m-%d"),
    }

    replace_placeholders(doc, mapping)

    filename = (
        f"{prefix}_{data.employee_name.replace(' ', '_')}_"
        f"{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    )

    file_path = os.path.join(GENERATED_DIR, filename)

    doc.save(file_path)

    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    review = ai_review_document(text, doc_type)

    return filename, review


# ---------------- ROUTES ----------------

@app.post("/generate-offer-letter")
def offer_letter(data: EmployeeData):

    f, r = generate_document(
        data,
        "offer_letter_template.docx",
        "OfferLetter",
        "Offer Letter"
    )

    return {"docx_file": f, "ai_review": r}


@app.post("/generate-appointment-letter")
def appointment_letter(data: EmployeeData):

    f, r = generate_document(
        data,
        "appointment_letter_template.docx",
        "AppointmentLetter",
        "Appointment Letter"
    )

    return {"docx_file": f, "ai_review": r}


@app.post("/generate-termination-letter")
def termination_letter(data: EmployeeData):

    f, r = generate_document(
        data,
        "termination_letter_template.docx",
        "TerminationLetter",
        "Termination Letter"
    )

    return {"docx_file": f, "ai_review": r}


@app.post("/generate-experience-letter")
def experience_letter(data: EmployeeData):

    f, r = generate_document(
        data,
        "experience_letter_template.docx",
        "ExperienceLetter",
        "Experience Letter"
    )

    return {"docx_file": f, "ai_review": r}


@app.get("/download/{filename}")
def download_file(filename: str):

    file_path = os.path.join(GENERATED_DIR, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )